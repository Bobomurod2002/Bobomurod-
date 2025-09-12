#1
import pandas as pd
from docx import Document

# Word faylini ochish
doc = Document(r"c:\Users\User\Documents\Python jadval.docx")

# Har bir jadvalni o‘qish (Wordda bir nechta jadval bo‘lishi mumkin)
table = doc.tables[0]  # birinchi jadvalni olish

# Jadvalni DataFrame ga aylantirish
data = []
keys = None
for i, row in enumerate(table.rows):
    text = [cell.text.strip() for cell in row.cells]
    if i == 0:
        keys = text  # birinchi qatordagi ustunlar nomi
    else:
        data.append(text)

df = pd.DataFrame(data, columns=keys)

# CSV faylga yozish
df.to_csv(r"c:\Users\User\Documents\stackoverflow_qa.csv", index=False, encoding="utf-8")

print("✅ Word fayli CSV formatga saqlandi!")

import pandas as pd

# CSV faylni o‘qish
#df = pd.read_csv(r "c:\Users\b.rahmatov\Desktop\11111.csv", parse_dates=['creationdate'])
df = pd.read_csv(r"c:\Users\User\Documents\stackoverflow_qa.csv", parse_dates=['creationdate'])

# 1. 2014-yildan oldin yaratilgan savollar
q1 = df[df['creationdate'] < '2014-01-01']

# 2. Score > 50
q2 = df[df['score'] > 50]

# 3. Score 50–100 oralig‘ida
q3 = df[(df['score'] >= 50) & (df['score'] <= 100)]

# 4. Scott Boston tomonidan javob berilgan savollar
q4 = df[df['ans_name'] == 'Scott Boston']

# 5. Quyidagi 5 ta foydalanuvchi javob bergan savollar
users = ['Scott Boston', 'Mike Pennington', 'unutbu', 'Demitri', 'doug']
q5 = df[df['ans_name'].isin(users)]

# 6. 2014-yil martdan oktabrgacha yaratilgan, javob bergan 'unutbu', score < 5
q6 = df[
    (df['creationdate'] >= '2014-03-01') & 
    (df['creationdate'] <= '2014-10-31') &
    (df['ans_name'] == 'unutbu') &
    (df['score'] < 5)
]

# 7. Score 5–10 oralig‘ida YOKI viewcount > 10,000
q7 = df[((df['score'] >= 5) & (df['score'] <= 10)) | (df['viewcount'] > 10000)]

# 8. Scott Boston tomonidan javob berilmagan savollar
q8 = df[df['ans_name'] != 'Scott Boston']


# Har bir natijani ko‘rish
print("1. 2014dan oldin:", q1.shape)
print("2. Score > 50:", q2.shape)
print("3. 50 <= Score <= 100:", q3.shape)
print("4. Scott Boston javob bergan:", q4.shape)
print("5. 5 ta foydalanuvchi javob bergan:", q5.shape)
print("6. 2014-03 dan 2014-10 gacha, unutbu, score<5:", q6.shape)
print("7. Score 5-10 yoki viewcount>10000:", q7.shape)
print("8. Scott Boston javob bermagan:", q8.shape)

#2
import pandas as pd
from docx import Document

# Word faylini ochish
doc = Document(r"c:\Users\User\Documents\Python jadval2.docx")

# Har bir jadvalni o‘qish (Wordda bir nechta jadval bo‘lishi mumkin)
table = doc.tables[0]  # birinchi jadvalni olish

# Jadvalni DataFrame ga aylantirish
data = []
keys = None
for i, row in enumerate(table.rows):
    text = [cell.text.strip() for cell in row.cells]
    if i == 0:
        keys = text  # birinchi qatordagi ustunlar nomi
    else:
        data.append(text)

df = pd.DataFrame(data, columns=keys)

# CSV faylga yozish
df.to_csv(r"c:\Users\User\Documents\titanic.csv", index=False, encoding="utf-8")

print("✅ Word fayli CSV formatga saqlandi!")

import pandas as pd

# Titanic datasetni yuklash
titanic_df = pd.read_csv("titanic.csv")

# Ustunlarning o‘zbekcha izohlari
columns_info_uz = {
    "PassengerId": "Har bir yo‘lovchining identifikatori (ID).",
    "Survived": "Yo‘lovchi tirik qolganmi yoki yo‘q (0 = Yo‘q, 1 = Ha).",
    "Pclass": "Chipta sinfi (1-sinf, 2-sinf yoki 3-sinf).",
    "Name": "Yo‘lovchining ismi.",
    "Sex": "Yo‘lovchining jinsi.",
    "Age": "Yo‘lovchining yoshi (yillarda).",
    "SibSp": "Bortda bo‘lgan aka-uka yoki turmush o‘rtoqlar soni.",
    "Parch": "Bortda bo‘lgan ota-ona yoki farzandlar soni.",
    "Ticket": "Chipta raqami.",
    "Fare": "Chipta narxi.",
    "Cabin": "Kajut (xona) raqami.",
    "Embarked": "Yo‘lovchi kemaga chiqqan port (C = Cherbourg, Q = Queenstown, S = Southampton)."
}

print("Titanic dataset ustunlari (O‘zbekcha izohlar):\n")
for col, desc in columns_info_uz.items():
    print(f"{col}: {desc}")

# Datasetdan kichik namuna
print("\nDatasetdan dastlabki 5 qator:")
print(titanic_df.head())


#3
import pandas as pd

# Titanic datasetni yuklash
titanic_df = pd.read_csv("titanic.csv")

# 1. 1-sinfdagi 20–30 yosh oralig‘idagi ayol yo‘lovchilar
female_class1_20_30 = titanic_df[
    (titanic_df['Sex'] == 'female') &
    (titanic_df['Pclass'] == 1) &
    (titanic_df['Age'].between(20, 30))
]

# 2. $100 dan ko‘p haq to‘lagan yo‘lovchilar
fare_over_100 = titanic_df[titanic_df['Fare'] > 100]

# 3. Yolg‘iz sayohat qilib omon qolgan yo‘lovchilar
survived_alone = titanic_df[
    (titanic_df['Survived'] == 1) &
    (titanic_df['SibSp'] == 0) &
    (titanic_df['Parch'] == 0)
]

# 4. 'C' portidan chiqqan va $50 dan ko‘p to‘lagan yo‘lovchilar
embarked_c_over_50 = titanic_df[
    (titanic_df['Embarked'] == 'C') &
    (titanic_df['Fare'] > 50)
]

# 5. Ham aka-uka/turmush o‘rtog‘i, ham ota-ona/farzandi bo‘lgan yo‘lovchilar
with_sibsp_parch = titanic_df[
    (titanic_df['SibSp'] > 0) &
    (titanic_df['Parch'] > 0)
]

# 6. 15 yoshdan kichik yoki teng va omon qolmagan yo‘lovchilar
age_15_or_younger_not_survived = titanic_df[
    (titanic_df['Age'] <= 15) &
    (titanic_df['Survived'] == 0)
]

# 7. Kajuti bor va $200 dan ko‘p to‘lagan yo‘lovchilar
cabin_and_fare_over_200 = titanic_df[
    titanic_df['Cabin'].notnull() &
    (titanic_df['Fare'] > 200)
]

# 8. PassengerId raqami toq bo‘lgan yo‘lovchilar
odd_passenger_id = titanic_df[titanic_df['PassengerId'] % 2 == 1]

# 9. Faqat bitta marta uchraydigan chipta raqamiga ega yo‘lovchilar
unique_ticket = titanic_df[
    titanic_df['Ticket'].isin(
        titanic_df['Ticket'].value_counts()[titanic_df['Ticket'].value_counts() == 1].index
    )
]

# 10. Ismida 'Miss' bo‘lgan va 1-sinfda bo‘lgan ayol yo‘lovchilar
miss_class1 = titanic_df[
    (titanic_df['Name'].str.contains("Miss")) &
    (titanic_df['Sex'] == 'female') &
    (titanic_df['Pclass'] == 1)
]

# Natija tekshirish (masalan, 1-variantdan dastlabki 5 qator)
print(miss_class1.head())



